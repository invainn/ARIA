import os
import json 

from flask import Flask, request

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

# Create the Flask app.
app = Flask(__name__)

# Generate documents and write them to S3
# Downloading documents will be done through a separate requests on Node.js
@app.route('/documents', methods=['POST'])
def generateDocuments():
    # Create document variables to hold processing status.
    # By default they are set to not being requested.
    # If the document is requested it will be processed successfully or unsuccessfully.
    # The string will be updated either way.
    adjudicationForm = 'document was not requested.'
    announcingSheet = 'document was not requested.'
    certificates = 'document was not requested.'
    masterJudgeRepertoire = 'document was not requested.'
    masterJudgeSchedule = 'document was not requested.'
    resultsSheet = 'document was not requested.'
    roomSchedules = 'document was not requested.'
    sessionAssignments = 'document was not requested.'
    sessionLabels = 'document was not requested.'
    teacherMaster = 'document was not requested.'

    # Read the incoming JSON data and store it in 'data'
    reqData = request.get_json()

    # Determine if the JSON object has data, if it does...process it.
    try:
        if (reqData is not None):
                # ----------------
                # Announcing Sheet
                # ----------------
            if (reqData['documents']['announcingSheet'] == True):
                announcingSheet = createAnnouncingSheet(reqData)
    except Exception as err:
        return f'Something Went Wrong \n Error Details Below \n, {err}'
    else:
        return f'''\n
                   01. Adjudication Form: {adjudicationForm} \n
                   02. Announcing Sheet: {announcingSheet} \n
                   03. Festival Certificates: {certificates} \n
                   04. Master Judge Repertoire: {masterJudgeRepertoire} \n
                   05. Master Judge Schedule: {masterJudgeSchedule} \n
                   06. Results Sheet: {resultsSheet} \n
                   07. Room Schedules: {roomSchedules} \n
                   08. Session Assignments: {sessionAssignments} \n
                   09. Session Labels: {sessionLabels} \n
                   10. Teacher Master: {teacherMaster} \n
        '''

def createAnnouncingSheet(data):
    document = Document()
    announcingSheet = data['announcingSheet']

    eventName = announcingSheet['eventName']
    documentTitle = announcingSheet['documentTitle']

    sessions = announcingSheet.get('sessions', 'Sessions not found')

    for session in sessions:
        day = session.get('day', 'Day not found')
        time = str(session['time'])
        sessionNumber = str(session['sessionNumber'])
        nameOfRoom = str(session['nameOfRoom'])
        location = day + ', ' + time + ', ' + sessionNumber + ', ' + nameOfRoom
        classType = session['classType']
        levels = session['levels']
        judges = session['judges']
        students = session['students']
        proctorName = session['proctorName']
        doorMonitorName = session['doorMonitorName']
        performanceOrder = session['performanceOrder']

        # Records Data For Table
        records = ()

        # Event Title
        document.add_heading(eventName, level=0)
        # Title Subheading
        document.add_heading(documentTitle, level=1)
        # Day, Time, Session Number, Room
        document.add_heading(location, level=2)
        # Add Performance Class Type
        document.add_heading(classType, level=1)
        # Add Students Levels
        document.add_heading(f'Levels: {levels}', level=2)

        # Add Judges
        for judge in judges:
            document.add_heading(f'Judge: {judge}', level=3)

        # Add Proctor
        document.add_heading(f'Proctor: {proctorName}', level=3)
        # Add Door Monitor
        document.add_heading(f'Door Monitor: {doorMonitorName}', level=3)
        # Performance Order
        document.add_heading(f'{performanceOrder}:\n', level=1)

        # Student Table
        table = document.add_table(rows=int((len(students) + 1)), cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        table.style = 'TableGrid'

        headerCells = table.rows[0].cells

        cell_headers = ['Order', 'Name', 'Level', 'Song I', 'Song II']

        cell_num = 0

        # Iterate through the headers and add them to the table
        # If it is a song, add another header with the composer
        for header in cell_headers:
            run = headerCells[cell_num].paragraphs[0].add_run(f'\n{header}\n')
            run.bold = True
            run.underline = True

            if 'Song' in header:
                run = headerCells[cell_num + 1].paragraphs[0].add_run(f'\n{header}\nComposer\n')
                run.bold = True
                run.underline = True

                cell_num = cell_num + 2
            else:
                cell_num = cell_num + 1


        for student in students:
            # Get Student Data
            student = tuple(student.values())
            records += (student,)

        # Add Student Data To Table
        i = 1
        for perf, lvl, songs in records:
            rowCells = table.rows[i].cells

            rowCells[0].text = str(i)
            rowCells[1].text = str(perf)
            rowCells[2].text = str(lvl)

            cellNumber = 3 

            # Iterate through each student's songs and insert them 
            for song in songs:
                rowRun = rowCells[cellNumber].paragraphs[0].add_run(str(song['songName'])).italic = True
                rowCells[cellNumber + 1].text = str(song['songComposer'])
                cellNumber = cellNumber + 2

            i += 1

        # Page break after each session object is processed.
        document.add_page_break()

        # # Write/Save Document
        document.save('./documents/announcing_sheet.docx')

        return 'document created successfully.'

if __name__ == "__main__":
    app.run(debug=True, port=4321) #run app in debug mode on port 4321
